from flask import Flask, request, jsonify, render_template, redirect, url_for, session,flash,make_response
from werkzeug.utils import secure_filename
from docx import Document
from io import BytesIO
from llm_api import call_llm_for_question,call_llm_for_feedback,call_llm_for_analysis,call_llm_for_ask,call_llm_for_realtime,evaluate_student_answer,call_llm_for_plan
from urllib.parse import quote
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import Workbook
from datetime import datetime
import sqlite3
import os
import re
import fitz
import zipfile

now =datetime.now().strftime('%Y-%m-%d %H:%M:%S')
app = Flask(__name__,template_folder="templates")
#用于Session管理
app.secret_key = 'secret_key'
UPLOAD_FOLDER = 'uploaded_knowledge'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = 'knowledge_files'

#判断文件是否合法
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

#用于解析文件类型
def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".txt":
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == ".docx":
        try:
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            return f"[DOCX 解析失败] {e}"
    elif ext == ".pdf":
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            return text
        except Exception as e:
            return f"[PDF 解析失败] {e}"
    else:
        return "[不支持的文件类型]"

#使用正则提取三段内容
def parse_generated_question(text):
    try:
        question = re.search(r"【题目】(.*?)【答案】", text, re.S).group(1).strip()
        answer = re.search(r"【答案】(.*?)【解析】", text, re.S).group(1).strip()
        explanation = re.search(r"【解析】(.*)", text, re.S).group(1).strip()
        return question, answer, explanation
    except:
        return text, "提取失败", "暂无解析"

#学习效果表
conn = sqlite3.connect('database.db')
cursor = conn.cursor()
cursor.execute('''
    CREATE TABLE IF NOT EXISTS student_learning_analysis (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        correct_trend TEXT,
        mastery_summary TEXT,
        top_wrong_knowledge TEXT,
        updated_by TEXT,
        updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        course_name TEXT,
        teacher_username TEXT
    )
''')
conn.commit()
conn.close()

#教学效率表
conn=sqlite3.connect('database.db')
cursor=conn.cursor()
cursor.execute('''
CREATE TABLE IF NOT EXISTS teaching_efficiency_analysis (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    course_name TEXT,
    teacher_username TEXT,
    plan_time_summary TEXT,
    practice_time_summary TEXT,
    suggestion_summary TEXT,
    updated_at TEXT
)
''')
conn.commit()
conn.close()

#日志表
conn = sqlite3.connect('database.db')
cursor = conn.cursor()
cursor.execute('''
CREATE TABLE IF NOT EXISTS system_logs (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    username TEXT,
    role TEXT,
    action TEXT,
    detail TEXT,
    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
)
''')
conn.commit()
conn.close()

#知识库表
conn = sqlite3.connect('database.db')
cursor = conn.cursor()
cursor.execute('''
    CREATE TABLE IF NOT EXISTS knowledge (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        course_name TEXT,
        teacher_username TEXT,
        filename TEXT,
        content TEXT,
        is_active INTEGER DEFAULT 1
    )
''')
conn.commit()
conn.close()

#课件表
conn = sqlite3.connect('database.db')
cursor = conn.cursor()
cursor.execute('''
CREATE TABLE IF NOT EXISTS teaching_plans (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    title TEXT NOT NULL,  -- 新增标题字段
    course_name TEXT NOT NULL,
    teacher_username TEXT NOT NULL,
    knowledge_ids TEXT,
    extra_request TEXT,
    plan_content TEXT,
    created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (course_name, teacher_username) REFERENCES courses(name, teacher_username)
)
''')
conn.commit()
conn.close()

#题库表
conn = sqlite3.connect('database.db')
cursor = conn.cursor()
cursor.execute('''
CREATE TABLE IF NOT EXISTS question_bank (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    question TEXT,
    answer TEXT,
    explanation TEXT,
    type TEXT,
    teacher_username TEXT,
    course_name TEXT
)
''')
conn.commit()
conn.close()

#做题记录
conn = sqlite3.connect('database.db')
cursor = conn.cursor()
cursor.execute('''
    CREATE TABLE IF NOT EXISTS answer_records (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    question_id INTEGER,
    question TEXT,
    student_answer TEXT,
    correct_answer TEXT,
    result INTEGER,
    student TEXT,
    type TEXT,
    teacher_username TEXT,
    course_name TEXT,
    feedback TEXT
);
''')
conn.commit()
conn.close()


#用户表
conn = sqlite3.connect('database.db')
cursor = conn.cursor()
cursor.execute('''
    CREATE TABLE if not EXISTS users (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    username TEXT UNIQUE,
    password TEXT,
    role TEXT,
    user_number TEXT
    )
''')
conn.commit()
conn.close()


#课程表
conn=sqlite3.connect('database.db')
cursor=conn.cursor()
cursor.execute('''
    CREATE TABLE IF NOT EXISTS courses (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL,
    teacher_username TEXT NOT NULL
);
''')
conn.commit()
conn.close()


#选课表
conn=sqlite3.connect('database.db')
cursor=conn.cursor()
cursor.execute('''
    CREATE TABLE IF NOT EXISTS student_courses (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    student_username TEXT NOT NULL,
    course_id INTEGER NOT NULL,
    FOREIGN KEY (student_username) REFERENCES users(username),
    FOREIGN KEY (course_id) REFERENCES courses(id)
    UNIQUE(student_username, course_id)  -- 禁止重复添加
);

''')
conn.commit()
conn.close()

@app.route('/')
def index():
    role = session.get('role')

    if role == 'student':
        return redirect('/student')
    elif role == 'teacher':
        return redirect('/teacher')
    elif role == 'admin':
        return redirect('/admin')
    else:
        return redirect('/login')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        uname = request.form['username']
        pwd = request.form['password']
        role = request.form['role']
        unumber= request.form.get('user_number')

        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        try:
            cursor.execute('INSERT INTO users (username, password, role,user_number) VALUES (?, ?, ?,?)', (uname, pwd, role,unumber))
            conn.commit()
        except:
            conn.close()
            return render_template('register.html', error='注册失败，用户名已存在')

        conn.close()
        return redirect('/login')

    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        uname = request.form['username']
        pwd = request.form['password']
        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        try:
            cursor.execute('SELECT role FROM users WHERE username = ? AND password = ?', (uname, pwd))
            row = cursor.fetchone()
        except sqlite3.OperationalError as e:
            conn.close()
            return f'数据库结构错误：{e}，请检查 users 表结构'
        conn.close()
        if row:
            session['username'] = uname
            session['role'] = row[0]
            return redirect('/')
        else:
            return render_template('login.html', error='用户名或密码错误')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('index'))

@app.route('/admin')
def admin_dashboard():
    if session.get('role') != 'admin':
        return '无权限访问'
    return render_template('admin.html', username=session.get('username'))

#用户管理
@app.route('/admin/user_manage')
def admin_user_manage():
    if session.get('role') != 'admin':
        return '无权限访问'
    keyword = request.args.get('keyword', '').strip()
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    if keyword:
        cursor.execute('''
            SELECT * FROM users
            WHERE username LIKE ? OR user_number LIKE ?
        ''', (f'%{keyword}%', f'%{keyword}%'))
    else:
        cursor.execute('SELECT * FROM users')
    users = cursor.fetchall()
    conn.close()
    return render_template('admin_user_manage.html', users=users, keyword=keyword)

#删除用户
@app.route('/admin/delete_user/<username>')
def delete_user(username):
    if session.get('role') != 'admin':
        return '无权限访问'
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("DELETE FROM users WHERE username = ?", (username,))
    cursor.execute('''
        INSERT INTO system_logs (username, role, action, detail)
        VALUES (?, ?, ?, ?)
    ''', (session['username'], 'admin', '删除用户', f'被删用户名: {username}'))
    conn.commit()
    conn.close()
    flash("用户已删除")
    return redirect('/admin/user_manage')

#教学资源管理
@app.route('/admin/teaching_resources')
def admin_teaching_resources():
    if session.get('role') != 'admin':
        return '无权限访问'
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    #获取所有课程列表
    cursor.execute('SELECT DISTINCT name FROM courses ORDER BY name')
    courses = [row[0] for row in cursor.fetchall()]
    #获取筛选参数
    selected_course = request.args.get('course')
    selected_teacher = request.args.get('teacher')
    teachers = []
    resources = []
    if selected_course:
        # 获取该课程下的教师列表
        cursor.execute('''
            SELECT DISTINCT teacher_username FROM courses 
            WHERE name = ? ORDER BY teacher_username
        ''', (selected_course,))
        teachers = [row[0] for row in cursor.fetchall()]
        if selected_teacher:
            # 获取该教师该课程下的所有资源
            # 1. 教学计划
            cursor.execute('''
                SELECT id, title, created_at FROM teaching_plans
                WHERE course_name = ? AND teacher_username = ?
                ORDER BY created_at DESC
            ''', (selected_course, selected_teacher))
            plans = cursor.fetchall()
            # 2. 知识库文件
            cursor.execute('''
                SELECT id, filename, created_at FROM knowledge
                WHERE course_name = ? AND teacher_username = ? AND is_active = 1
                ORDER BY created_at DESC
            ''', (selected_course, selected_teacher))
            knowledge_files = cursor.fetchall()
            # 3. 题目数量统计
            cursor.execute('''
                SELECT type, COUNT(*) FROM question_bank
                WHERE course_name = ? AND teacher_username = ?
                GROUP BY type
            ''', (selected_course, selected_teacher))
            question_stats = cursor.fetchall()
            resources = {
                'plans': plans,
                'knowledge_files': knowledge_files,
                'question_stats': question_stats
            }
    conn.close()
    return render_template('admin_teaching_resources.html',
                            courses=courses,
                            teachers=teachers,
                            selected_course=selected_course,
                            selected_teacher=selected_teacher,
                            resources=resources)

#导出资料
@app.route('/admin/export_resources')
def admin_export_resources():
    if session.get('role') != 'admin':
        return '无权限访问'
    course = request.args.get('course')
    teacher = request.args.get('teacher')
    resource_type = request.args.get('type')  # 'plans', 'knowledge' 或 'questions'
    if not all([course, teacher, resource_type]):
        return '参数不完整'
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"{course}_{teacher}_{resource_type}_{timestamp}"
    safe_filename = quote(filename)
    if resource_type == 'plans':
        # 导出教学计划为Word
        cursor.execute('''
            SELECT title, plan_content FROM teaching_plans
            WHERE course_name = ? AND teacher_username = ?
            ORDER BY created_at DESC
        ''', (course, teacher))
        plans = cursor.fetchall()
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.add_heading(f'{course} 教学计划', level=1)
        doc.add_paragraph(f'教师: {teacher}')
        doc.add_paragraph(f'导出时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        for title, content in plans:
            doc.add_heading(title or "教学计划", level=2)
            doc.add_paragraph(content)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        response = make_response(buffer.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{safe_filename}.docx'
    elif resource_type == 'knowledge':
        # 知识库保持ZIP压缩格式
        cursor.execute('''
            SELECT filename, content FROM knowledge
            WHERE course_name = ? AND teacher_username = ? AND is_active = 1
            ORDER BY created_at DESC
        ''', (course, teacher))
        knowledge_files = cursor.fetchall()
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for filename, content in knowledge_files:
                zip_file.writestr(filename, content)
        response = make_response(zip_buffer.getvalue())
        response.headers['Content-Type'] = 'application/zip'
        response.headers['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{safe_filename}.zip'
    elif resource_type == 'questions':
        q_type = request.args.get('q_type')  # 可选
        base_sql = '''
            SELECT question, answer, type FROM question_bank
            WHERE course_name = ? AND teacher_username = ?
        '''
        params = [course, teacher]
        if q_type:
            base_sql += ' AND type = ?'
            params.append(q_type)
        base_sql += ' ORDER BY type, id'
        cursor.execute(base_sql, params)
        questions = cursor.fetchall()
        # 用openpyxl写Excel（更专业）
        wb = Workbook()
        ws = wb.active
        ws.title = "题目列表"
        # 写入标题
        ws.append(['题干', '答案', '题型'])
        # 写入每条数据
        for q in questions:
            ws.append(list(q))
        # 保存为字节流
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        response = make_response(buffer.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        response.headers['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{safe_filename}.xlsx'
    else:
        conn.close()
        return '无效的资源类型'
    cursor.execute('''
        INSERT INTO system_logs (username, role, action, detail)
        VALUES (?, ?, ?, ?)
    ''', (session['username'], 'admin', '导出教学资源',
          f'课程: {course}, 教师: {teacher}, 类型: {resource_type}'))
    conn.commit()
    conn.close()
    return response

def get_top_action(stats_dict):
    return max(stats_dict.items(), key=lambda x: x[1])[0] if stats_dict else "暂无数据"

#管理端大屏
@app.route('/admin/dashboard')
def admin_dashboard_stats():
    if session.get('role') != 'admin':
        return '无权限访问'
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 教师上传 + 出题
    cursor.execute('SELECT COUNT(*) FROM knowledge WHERE DATE(created_at, "localtime") = DATE("now", "localtime")')
    teacher_upload_today = cursor.fetchone()[0]
    cursor.execute('SELECT COUNT(*) FROM question_bank WHERE DATE(created_at, "localtime") = DATE("now", "localtime")')
    teacher_question_today = cursor.fetchone()[0]
    teacher_today_total = teacher_upload_today + teacher_question_today
    # 教师本周活跃课程
    cursor.execute('''
        SELECT course_name, COUNT(*) FROM question_bank
        WHERE DATE(created_at, "localtime") >= DATE("now", "-7 day", "localtime")
        GROUP BY course_name ORDER BY COUNT(*) DESC LIMIT 5
    ''')
    teacher_week_courses = cursor.fetchall()
    #今日使用统计函数
    def count_action_today(role, action):
        cursor.execute('''
            SELECT COUNT(*) FROM system_logs
            WHERE role = ? AND action = ? AND DATE(timestamp, 'localtime') = DATE('now', 'localtime')
        ''', (role, action))
        return cursor.fetchone()[0]
    #一周使用统计函数
    def count_action_weekly(role, action):
        cursor.execute('''
            SELECT COUNT(*) FROM system_logs
            WHERE role = ? AND action = ? AND DATE(timestamp, 'localtime') >= DATE('now', '-7 day', 'localtime')
        ''', (role, action))
        return cursor.fetchone()[0]
    #学生功能使用统计（今日）
    student_usage_stats = {
        "添加课程": count_action_today("student", "添加课程"),
        "退选课程": count_action_today("student", "退选课程"),
        "开始练习": count_action_today("student", "开始练习"),
        "查看记录": count_action_today("student", "查看记录"),
        "开始实时练习": count_action_today("student", "开始实时练习"),
        "学习助手问答": count_action_today("student", "学习助手问答")
    }
    student_today_total = sum(student_usage_stats.values())
    # 学生功能使用统计（本周）
    student_usage_weekly = {
        k: count_action_weekly("student", k) for k in student_usage_stats.keys()
    }
    student_week_total = sum(student_usage_weekly.values())
    # 学生今日做题数
    cursor.execute('''
        SELECT COUNT(*) FROM answer_records
        WHERE DATE(created_at, 'localtime') = DATE('now', 'localtime')
    ''')
    student_answer_today = cursor.fetchone()[0]
    # 学生本周做题数
    cursor.execute('''
        SELECT COUNT(*) FROM answer_records
        WHERE DATE(created_at, 'localtime') >= DATE('now', '-7 day', 'localtime')
    ''')
    student_answer_week = cursor.fetchone()[0]
    # 教师功能使用统计（今日）
    teacher_usage_stats = {
        "上传知识库": count_action_today("teacher", "上传知识库"),
        "删除知识库": count_action_today("teacher", "删除知识库"),
        "生成教学设计": count_action_today("teacher", "生成教学设计"),
        "生成题目": count_action_today("teacher", "生成题目"),
        "查看题库": count_action_today("teacher", "查看题库"),
        "添加课程": count_action_today("teacher", "添加课程"),
        "删除课程": count_action_today("teacher", "删除课程"),
        "添加学生": count_action_today("teacher", "添加学生"),
        "删除学生": count_action_today("teacher", "删除学生"),
        "编辑题库": count_action_today("teacher", "编辑题库"),
        "删除题库": count_action_today("teacher", "删除题库"),
        "修正教学设计": count_action_today("teacher", "修正教学设计"),
        "删除教学设计": count_action_today("teacher", "删除教学设计")
    }
    teacher_today_total = sum(teacher_usage_stats.values())
    # 教师功能使用统计（本周）
    teacher_usage_weekly = {
        k: count_action_weekly("teacher", k) for k in teacher_usage_stats.keys()
    }
    teacher_week_total = sum(teacher_usage_weekly.values())
    top_teacher_action_today = get_top_action(teacher_usage_stats)
    top_teacher_action_week = get_top_action(teacher_usage_weekly)
    top_student_action_today = get_top_action(student_usage_stats)
    top_student_action_week = get_top_action(student_usage_weekly)
    # 学生正确率
    cursor.execute('''
        SELECT AVG(result) FROM answer_records
        WHERE DATE(created_at, 'localtime') >= DATE('now', '-7 day', 'localtime')
    ''')
    avg_score = cursor.fetchone()[0]
    avg_correct_rate = f"{avg_score:.1f}%" if avg_score is not None else "暂无数据"
    # 高频错题
    cursor.execute('''
        SELECT question, COUNT(*) FROM answer_records
        WHERE DATE(created_at, 'localtime') >= DATE('now', '-7 day', 'localtime') AND result < 90
        GROUP BY question ORDER BY COUNT(*) DESC LIMIT 5
    ''')
    # 学生学习分析结果（只取最近一次）
    cursor.execute('''
    SELECT course_name, teacher_username, correct_trend, mastery_summary, top_wrong_knowledge, updated_at
    FROM student_learning_analysis
    ORDER BY updated_at DESC LIMIT 1
''')
    row = cursor.fetchone()
    learning_analysis = {
        "course_name": row[0],
        "teacher_username": row[1],
        "correct_trend": row[2],
        "mastery_summary": row[3],
        "top_wrong_knowledge": row[4],
        "updated_at": row[5]
    } if row else None
    cursor.execute('''
    SELECT course_name, teacher_username, plan_time_summary, practice_time_summary, suggestion_summary, updated_at
    FROM teaching_efficiency_analysis
    ORDER BY updated_at DESC LIMIT 1
    ''')
    row2 = cursor.fetchone()
    efficiency_analysis = {
        "course_name": row2[0],
        "teacher_username": row2[1],
        "plan_time_summary": row2[2],
        "practice_time_summary": row2[3],
        "suggestion_summary": row2[4],
        "updated_at": row2[5]
    } if row2 else None
    conn.close()
    print("top_teacher_action_today: ", top_teacher_action_today)
    return render_template('dashboard.html',
                           teacher_upload_today=teacher_upload_today,
                           teacher_question_today=teacher_question_today,
                           teacher_today_total=teacher_today_total,
                           teacher_week_courses=teacher_week_courses,
                           teacher_usage_stats=teacher_usage_stats,
                           teacher_today_total_usage=teacher_today_total,
                           teacher_usage_weekly=teacher_usage_weekly,
                           teacher_week_total=teacher_week_total,
                           student_usage_stats=student_usage_stats,
                           student_usage_weekly=student_usage_weekly,
                           student_today_total=student_today_total,
                           student_week_total=student_week_total,
                           student_answer_today=student_answer_today,
                           student_answer_week=student_answer_week,
                           learning_analysis=learning_analysis,
                           correct_rate=avg_correct_rate,
                           top_teacher_action_today=top_teacher_action_today,
                           top_teacher_action_week=top_teacher_action_week,
                           top_student_action_today=top_student_action_today,
                           top_student_action_week=top_student_action_week,
                           efficiency_analysis=efficiency_analysis
                           )
#课程筛选
@app.route('/admin/search_courses')
def admin_search_courses():
    keyword = request.args.get('keyword', '').strip()
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('''
        SELECT DISTINCT name, teacher_username FROM courses
        WHERE name LIKE ?
        ORDER BY name LIMIT 10
    ''', (f'%{keyword}%',))
    results = [{"course_name": r[0], "teacher_username": r[1]} for r in cursor.fetchall()]
    conn.close()
    return jsonify(results)

#学生学习效果分析
@app.route('/admin/update_learning_analysis', methods=['POST'])
def update_learning_analysis():
    if session.get('role') != 'admin':
        return '无权限访问'
    selected = request.form.get('course_teacher')
    if not selected or '|' not in selected:
        return "请选择课程和教师"
    course_name, teacher_username = selected.split('|')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('''
        SELECT course_name, question, result FROM answer_records
        WHERE course_name = ? AND teacher_username = ?
        AND DATE(created_at, 'localtime') >= DATE('now', '-14 day', 'localtime')
    ''', (course_name, teacher_username))
    records = cursor.fetchall()
    conn.close()
    if not records:
        flash("该课程暂无足够答题记录")
        return redirect('/admin/dashboard')
    plain_text = "\n".join([f"课程：{r[0]}\n题目：{r[1]}\n得分：{r[2]}\n---" for r in records])
    prompt = f"""你是一个教学评估专家，请根据以下答题数据，输出三段独立分析（分别为：正确率趋势、掌握情况、高频错误知识点）。
{plain_text}"""
    result = call_llm_for_analysis(prompt)
    parts = result.strip().split('\n\n')
    correct_trend = parts[0].strip() if len(parts) > 0 else '暂无数据'
    mastery_summary = parts[1].strip() if len(parts) > 1 else '暂无数据'
    top_wrong_knowledge = parts[2].strip() if len(parts) > 2 else '暂无数据'
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cursor.execute('''
        INSERT INTO student_learning_analysis
        (course_name, teacher_username, correct_trend, mastery_summary, top_wrong_knowledge, updated_at)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (course_name, teacher_username, correct_trend, mastery_summary, top_wrong_knowledge,now))
    conn.commit()
    conn.close()
    flash("更新成功")
    return redirect('/admin/dashboard')

#教学效率指标分析
@app.route('/admin/update_efficiency', methods=['POST'])
def update_efficiency():
    if session.get('role') != 'admin':
        return '无权限访问'
    value = request.form.get('eff_course_teacher')
    if not value or '|' not in value:
        return '课程选择无效'
    course_name, teacher_username = value.split('|')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 获取教学计划时间数据
    cursor.execute('''
        SELECT total_revision_seconds FROM teaching_plans
        WHERE course_name = ? AND teacher_username = ?
    ''', (course_name, teacher_username))
    plan_times = [r[0] or 0 for r in cursor.fetchall()]
    plan_total = sum(plan_times)
    plan_avg = plan_total // len(plan_times) if plan_times else 0
    plan_summary = f"总耗时 {plan_total} 秒，平均耗时 {plan_avg} 秒"
    # 获取题目修正数据
    cursor.execute('''
        SELECT type, total_revision_seconds FROM question_bank
        WHERE course_name = ? AND teacher_username = ?
    ''', (course_name, teacher_username))
    type_time_map = {}
    for qtype, seconds in cursor.fetchall():
        if qtype not in type_time_map:
            type_time_map[qtype] = []
        type_time_map[qtype].append(seconds or 0)
    all_total = sum(sum(lst) for lst in type_time_map.values())
    practice_lines = [f"总耗时 {all_total} 秒"]
    for qtype, lst in type_time_map.items():
        total = sum(lst)
        avg = total // len(lst) if lst else 0
        practice_lines.append(f"{qtype}题型：总 {total} 秒，平均 {avg} 秒")
    practice_summary = '；'.join(practice_lines)
    # 分析课程优化方向
    cursor.execute('''
        SELECT result, question, student_answer, correct_answer, type
        FROM answer_records
        WHERE course_name = ? AND teacher_username = ?
    ''', (course_name, teacher_username))
    records = cursor.fetchall()
    all_scores = [r[0] for r in records if r[0] is not None]
    type_scores = {}
    for score, _, _, _, qtype in records:
        if score is not None:
            type_scores.setdefault(qtype, []).append(score)
    avg_all = round(sum(all_scores) / len(all_scores), 1) if all_scores else 0
    type_lines = [f"整体平均分：{avg_all}"]
    for qtype, lst in type_scores.items():
        avg = round(sum(lst) / len(lst), 1) if lst else 0
        type_lines.append(f"{qtype}题型平均分：{avg}")
    # 构造提示文本送入大模型
    prompt = f"""请基于以下学生答题数据，总结出当前课程《{course_name}》的教学优化方向，建议风格为精准简练，如输出若干关键词 + 1-2句总结：
- {len(records)}条记录
- {'；'.join(type_lines)}
- 答题错误示例（前5条）：
""" + "\n".join(
        [f"题目：{q} 学生答：{sa} 正确答：{ca}" for _, q, sa, ca, _ in records[:5]]
    )
    suggestion = call_llm_for_plan(prompt)  # 重用原有模型调用函数
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    cursor.execute('''
        INSERT INTO teaching_efficiency_analysis (course_name, teacher_username, plan_time_summary, practice_time_summary, suggestion_summary, updated_at)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (course_name, teacher_username, plan_summary, practice_summary, suggestion, now))
    conn.commit()
    conn.close()
    return redirect('/admin/dashboard')

@app.route('/teacher')
def teacher_home():
    if session.get('role') != 'teacher':
        return '无权限访问'
    return render_template('teacher.html', username=session.get('username'))

@app.route('/get_knowledge_by_course')
def get_knowledge_by_course():
    course_name = request.args.get('course_name')
    teacher = session.get('username')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('''
        SELECT id, filename FROM knowledge
        WHERE course_name = ? AND teacher_username = ? AND is_active = 1
    ''', (course_name, teacher))
    rows = cursor.fetchall()
    conn.close()
    return jsonify([{'id': row[0], 'filename': row[1]} for row in rows])

#生成教学计划
@app.route('/teacher/generate_plan', methods=['GET', 'POST'])
def generate_plan():
    if session.get('role') != 'teacher':
        return '无权限访问'
    teacher = session['username']
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("SELECT id, name, teacher_username FROM courses WHERE teacher_username = ?", (teacher,))
    courses = cursor.fetchall()
    if request.method == 'GET':
        return render_template('generate_plan.html', courses=courses)
    course_id = request.form.get('course_id')
    knowledge_id = request.form.get('knowledge_id')
    extra = request.form.get('extra_request') or ""
    # 获取课程名
    cursor.execute("SELECT name FROM courses WHERE id = ? AND teacher_username = ?", (course_id, teacher))
    course_row = cursor.fetchone()
    if not course_row:
        conn.close()
        return '非法课程'
    course_name = course_row[0]
    # 获取知识库内容
    if knowledge_id == 'all':
        cursor.execute("SELECT content, id FROM knowledge WHERE course_name = ? AND teacher_username = ? AND is_active = 1", (course_name, teacher))
    else:
        cursor.execute("SELECT content, id FROM knowledge WHERE id = ? AND teacher_username = ? AND is_active = 1", (knowledge_id, teacher))
    knowledge_entries = cursor.fetchall()
    if not knowledge_entries:
        conn.close()
        return '未找到知识库'
    # 拼接知识文本与ID列表
    full_knowledge = "\n".join([k[0] for k in knowledge_entries])
    all_ids = ",".join(str(k[1]) for k in knowledge_entries)
    # 构造 Prompt
    prompt = f"""你是一位高校教师备课助手，请结合以下课程资料与要求，自动生成教学计划设计：
1. 教学内容概述（300字左右）
2. 实训安排建议（列出实训主题 + 内容要点）
3. 教学时间分布（如按章节或主题分配周次）
【课程名称】：{course_name}
【教师补充要求】：{extra if extra.strip() else '无'}
【教学知识内容】：\n{full_knowledge}
"""
    plan_title = request.form.get('plan_title')
    result = call_llm_for_plan(prompt)
    now =datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    # 写入 teaching_plans 表
    cursor.execute('''
        INSERT INTO teaching_plans (title,course_name, teacher_username, knowledge_ids, extra_request, plan_content,created_at,last_revision_start)
        VALUES (?,?, ?, ?, ?, ?,?,?)
    ''', (plan_title, course_name, teacher, all_ids, extra, result,now,now))
    plan_id = cursor.lastrowid
    start_time = now
    cursor.execute('''
    UPDATE teaching_plans SET last_revision_start = ?
    WHERE id = ?
''', (start_time, plan_id))
    # 写入日志
    cursor.execute('''
        INSERT INTO system_logs (username, role, action, detail)
        VALUES (?, ?, ?, ?)
    ''', (teacher, 'teacher', '生成教学设计', f'课程: {course_name}, 知识库: {all_ids}'))
    conn.commit()
    conn.close()
    return render_template('generate_plan_result.html', result=result, plan_id=plan_id)

#查看教学计划
@app.route('/teacher/plan_list')
def teacher_plan_list():
    if session.get('role') != 'teacher':
        return '无权限访问'
    teacher = session['username']
    search_query = request.args.get('search', '').strip()
    selected_course = request.args.get('course', '').strip()
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 获取该教师所有课程名（供下拉选择）
    cursor.execute("SELECT DISTINCT course_name FROM teaching_plans WHERE teacher_username = ?", (teacher,))
    course_names = [row[0] for row in cursor.fetchall()]
    # 构建 SQL 筛选
    sql = "SELECT id, title, course_name, created_at, total_revision_seconds FROM teaching_plans WHERE teacher_username = ?"
    params = [teacher]
    if selected_course:
        sql += " AND course_name = ?"
        params.append(selected_course)
    if search_query:
        sql += " AND title LIKE ?"
        params.append(f"%{search_query}%")
    sql += " ORDER BY created_at DESC"
    cursor.execute(sql, params)
    plans = cursor.fetchall()
    conn.commit()
    conn.close()
    return render_template('plan_list.html',
                           plans=plans,
                           course_names=course_names,
                           selected_course=selected_course,
                           search_query=search_query)

#编辑教学计划
@app.route('/teacher/edit_plan/<int:plan_id>')
def edit_plan(plan_id):
    if session.get('role') != 'teacher':
        return '无权限访问'
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('''
        SELECT plan_content FROM teaching_plans WHERE id = ? AND teacher_username = ?
    ''', (plan_id, session['username']))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return '未找到计划'
    # 更新开始时间为当前
    cursor.execute('''
        UPDATE teaching_plans SET last_revision_start = datetime('now', 'localtime') WHERE id = ?
    ''', (plan_id,))
    conn.commit()
    conn.close()
    return render_template('generate_plan_result.html', result=row[0], plan_id=plan_id)

#保存教学计划修改
@app.route('/teacher/save_plan_revision', methods=['POST'])
def save_plan_revision():
    if session.get('role') != 'teacher':
        return '无权限访问'
    plan_id = int(request.form.get('plan_id'))
    print("接收到的plan_id", plan_id)
    content = request.form.get('plan_content')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 获取上次修正时间（可空）
    cursor.execute("SELECT last_revision_start, total_revision_seconds FROM teaching_plans WHERE id = ?", (plan_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return '教学计划不存在'
    last_start = row[0]
    total = row[1] or 0
    elapsed = 0
    if last_start:
        from datetime import datetime
        fmt = "%Y-%m-%d %H:%M:%S"
        start_time = datetime.strptime(last_start, fmt)
        elapsed = int((datetime.now() - start_time).total_seconds())
    new_total = total + elapsed
    cursor.execute('''
        UPDATE teaching_plans SET plan_content = ?, total_revision_seconds = ?, last_revision_start = NULL
        WHERE id = ?
    ''', (content, new_total, plan_id))
    cursor.execute('''
        INSERT INTO system_logs (username, role, action, detail)
        VALUES (?, ?, ?, ?)
    ''', (session['username'], 'teacher', '修正教学设计', f'教学计划ID: {plan_id}, 用时: {elapsed}秒'))
    conn.commit()
    conn.close()
    return redirect('/teacher/plan_list')

#删除教学计划
@app.route('/teacher/delete_plan/<int:plan_id>', methods=['POST'])
def delete_plan(plan_id):
    if session.get('role') != 'teacher':
        return '无权限访问'
    teacher = session['username']
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('''
        DELETE FROM teaching_plans
        WHERE id = ? AND teacher_username = ?
    ''', (plan_id, teacher))
    cursor.execute('''
        INSERT INTO system_logs (username, role, action, detail)
        VALUES (?, ?, ?, ?)
    ''', (teacher, 'teacher', '删除教学设计', f'教学计划ID: {plan_id}'))
    conn.commit()
    conn.close()
    return redirect('/teacher/plan_list')

#生成题目
@app.route('/teacher/generate_question', methods=['GET', 'POST'])
def generate_question():
    if session.get('role') != 'teacher':
        return '无权限访问'
    teacher = session.get('username')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    if request.method == 'GET':
        question_type = request.form.get('type')
        cursor.execute("SELECT id, filename, course_name FROM knowledge WHERE teacher_username = ? AND is_active = 1", (teacher,))
        knowledge_list = cursor.fetchall()
        cursor.execute("SELECT name FROM courses WHERE teacher_username = ?", (teacher,))
        courses = [c[0] for c in cursor.fetchall()]
        conn.close()
        return render_template('generate_question.html', knowledge_list=knowledge_list, courses=courses)
    elif request.method == 'POST':
        question_type = request.form.get('type')
        knowledge_id = request.form.get('knowledge_id')
        selected_course = request.form.get('course')
        if knowledge_id == 'all':
            cursor.execute("SELECT content FROM knowledge")
        else:
            cursor.execute("SELECT content FROM knowledge WHERE id = ?", (knowledge_id,))
        knowledge_entries = cursor.fetchall()
        conn.close()
        full_knowledge = "\n".join([entry[0] for entry in knowledge_entries])
        prompt = f"""
你是一个专业教育系统中的AI出题助手，请根据“知识内容”，生成一道“{question_type}”题型的题目。你必须严格输出下列三个部分，且每部分开头必须加标识：

【题目】
...

【答案】
...

【解析】
...

注意；choice代表选择题，blank代表填空题，short代表简答题，code代表编程题，禁止无视题型乱生成题目!!!!!
注意：如果生成的是选择题，答案只需要带有选项的编号即可，禁止带有选项的具体内容，此则注意仅适用于选择题!!!!!
注意：如果生成的填空题，题目需要用下划线来挖空，答案即为该空的具体内容，此则注意仅适用于填空题!!!!!
注意：如果生成是简答题，则答案为一段文字，此则注意仅适用于简答题!!!!!
注意：如果生成是编程题，则答案为一段代码，此则注意仅适用于编程题!!!!!
注意：上述5则注意连同此条都必须同时遵守!!!!!
【知识内容如下】
{full_knowledge}
"""
        llm_response = call_llm_for_question(prompt,question_type)
        selected_course = request.form.get("course_name")
        question, answer, explanation = parse_generated_question(llm_response)
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn = sqlite3.connect('database.db')
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO question_bank (question, answer, explanation, type, teacher_username, course_name,created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (question, answer, explanation, question_type, teacher, selected_course,now))
        cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (teacher, 'teacher', '生成题目', f'{selected_course}：{question_type}'))
        conn.commit()
        conn.close()
        return redirect('/teacher/generate_question')

#查看题库
@app.route('/teacher/question_bank', methods=['GET', 'POST'])
def view_question_bank():
    teacher = session.get('username')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 获取筛选参数
    selected_course = request.args.get('course', '')
    selected_type = request.args.get('type', '')
    keyword = request.args.get('keyword', '')
    query = '''
        SELECT id, question, answer, explanation, type, course_name, created_at, total_revision_seconds
        FROM question_bank
        WHERE teacher_username = ?
    '''
    params = [teacher]
    if selected_course:
        query += " AND course_name = ?"
        params.append(selected_course)
    if selected_type:
        query += " AND type = ?"
        params.append(selected_type)
    if keyword:
        query += " AND (question LIKE ? OR answer LIKE ? OR explanation LIKE ?)"
        kw_like = f'%{keyword}%'
        params.extend([kw_like, kw_like, kw_like])
    query += " ORDER BY created_at DESC"
    cursor.execute(query, tuple(params))
    results = cursor.fetchall()
    # 获取课程选项和题型选项
    cursor.execute("SELECT DISTINCT course_name FROM question_bank WHERE teacher_username = ?", (teacher,))
    courses = [row[0] for row in cursor.fetchall()]
    cursor.execute("SELECT DISTINCT type FROM question_bank WHERE teacher_username = ?", (teacher,))
    types = [row[0] for row in cursor.fetchall()]
    conn.close()
    return render_template('question_bank.html',
                           questions=results,
                           courses=courses,
                           types=types,
                           selected_course=selected_course,
                           selected_type=selected_type,
                           keyword=keyword)

#编辑题目
@app.route('/teacher/edit_question/<int:question_id>')
def edit_question(question_id):
    if session.get('role') != 'teacher':
        return '无权限访问'
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('SELECT id, question, answer, explanation FROM question_bank WHERE id = ?', (question_id,))
    row = cursor.fetchone()
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    # 设置 last_revision_start 为当前时间
    cursor.execute("UPDATE question_bank SET last_revision_start = ? WHERE id = ?", (now, question_id))
    conn.commit()
    conn.close()
    if not row:
        return '题目不存在'
    return render_template('edit_question.html', question={
        'id': row[0],
        'question': row[1],
        'answer': row[2],
        'explanation': row[3]
    }, from_url=request.referrer or '/teacher/question_bank')


@app.route('/teacher/update_question', methods=['POST'])
def update_question():
    if session.get('role') != 'teacher':
        return '无权限访问'
    try:
        qid = int(request.form.get('id'))
    except (TypeError, ValueError):
        return '题目 ID 无效'
    question = request.form.get('question')
    answer = request.form.get('answer')
    explanation = request.form.get('explanation')
    from_url = request.form.get('from_url') or '/teacher/question_bank'
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 获取上次修正时间与总耗时
    cursor.execute("SELECT last_revision_start, total_revision_seconds FROM question_bank WHERE id = ?", (qid,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return '题目不存在'
    last_start = row[0]
    total = row[1] or 0
    elapsed = 0
    if last_start:
        fmt = "%Y-%m-%d %H:%M:%S"
        start_time =datetime.strptime(last_start, fmt)
        elapsed = int((datetime.now() - start_time).total_seconds())
    new_total = total + elapsed
    cursor.execute('''
        UPDATE question_bank SET question = ?, answer = ?, explanation = ?, total_revision_seconds = ?, last_revision_start = NULL
        WHERE id = ?
    ''', (question, answer, explanation, new_total, qid))
    cursor.execute('''
        INSERT INTO system_logs (username, role, action, detail)
        VALUES (?, ?, ?, ?)
    ''', (session['username'], 'teacher', '编辑题库', f'题目ID: {qid}, 用时: {elapsed}秒'))
    conn.commit()
    conn.close()
    return redirect(from_url)

#保存题目修改
@app.route('/teacher/save_question_revision', methods=['POST'])
def save_question_revision():
    if session.get('role') != 'teacher':
        return '无权限访问'
    question_id = int(request.form.get('question_id'))
    new_question = request.form.get('question')
    new_answer = request.form.get('answer')
    new_explanation = request.form.get('explanation')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('SELECT last_revision_start, total_revision_seconds FROM question_bank WHERE id = ?', (question_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return '题目不存在'
    last_start = row[0]
    total = row[1] or 0
    elapsed = 0
    if last_start:
        from datetime import datetime
        fmt = "%Y-%m-%d %H:%M:%S"
        start_time = datetime.strptime(last_start, fmt)
        elapsed = int((datetime.now() - start_time).total_seconds())
    new_total = total + elapsed
    cursor.execute('''
        UPDATE question_bank
        SET question = ?, answer = ?, explanation = ?, total_revision_seconds = ?, last_revision_start = NULL
        WHERE id = ?
    ''', (new_question, new_answer, new_explanation, new_total, question_id))
    cursor.execute('''
        INSERT INTO system_logs (username, role, action, detail)
        VALUES (?, ?, ?, ?)
    ''', (session['username'], 'teacher', '编辑题库', f'题目ID: {question_id}, 用时: {elapsed}秒'))
    conn.commit()
    conn.close()
    return redirect('/teacher/question_bank')

#删除题目
@app.route('/teacher/delete_question/<int:question_id>')
def delete_question(question_id):
    if session.get('role') != 'teacher':
        return '无权限访问'
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("DELETE FROM question_bank WHERE id = ?", (question_id,))
    cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (session['username'], 'teacher', '删除题库', f'题目ID: {question_id}'))
    conn.commit()
    conn.close()
    return redirect('/teacher/question_bank')

#查看知识库
@app.route('/teacher/knowledge', methods=['GET', 'POST'])
def manage_knowledge():
    if session.get('role') != 'teacher':
        return '无权限访问'
    teacher = session.get('username')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 获取教师教授的课程
    cursor.execute("SELECT name FROM courses WHERE teacher_username = ?", (teacher,))
    courses = [row[0] for row in cursor.fetchall()]
    # 处理上传逻辑
    if request.method == 'POST':
        course_name = request.form.get('course_name')
        file = request.files.get('file')
        if not course_name or not file or file.filename == '':
            flash("请选择课程并上传有效文件")
        elif allowed_file(file.filename):
            filename = secure_filename(file.filename)
            save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            file.save(save_path)
            content = extract_text(save_path)
            now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            cursor.execute('''
                INSERT INTO knowledge (filename, content, course_name, teacher_username,is_active, created_at)
                VALUES (?, ?, ?, ?,1,?)
            ''', (filename, content, course_name, teacher,now))
            cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (teacher, 'teacher', '上传知识库', f'{course_name}：{filename}'))
            conn.commit()
            flash("上传成功")
        else:
            flash("文件格式不支持（仅支持 PDF、TXT、DOCX）")
    # 获取该教师上传的知识库，按课程分类
    cursor.execute('''
        SELECT id, filename, course_name FROM knowledge
        WHERE teacher_username = ? AND is_active = 1
        ORDER BY course_name
    ''', (teacher,))
    records = cursor.fetchall()
    conn.close()
    # 分类整理
    knowledge_by_course = {}
    for rid, fname, cname in records:
        knowledge_by_course.setdefault(cname, []).append({'id': rid, 'filename': fname})

    return render_template('knowledge_manage.html',
                           courses=courses,
                           knowledge_by_course=knowledge_by_course)

@app.route('/teacher/delete_knowledge/<int:kid>', methods=['POST'])
def delete_knowledge(kid):
    if session.get('role') != 'teacher':
        return '无权限访问'
    teacher = session.get('username')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('''
        DELETE FROM knowledge
        WHERE id = ? AND teacher_username = ?
    ''', (kid, teacher))
    cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (teacher, 'teacher', '删除知识库', f'知识库ID: {kid}'))
    conn.commit()
    conn.close()
    flash("已删除知识库文件")
    return redirect('/teacher/knowledge')

@app.route('/teacher/student_records')
def student_records():
    if session.get('role') != 'teacher':
        return '无权限访问'
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('SELECT student, question, student_answer, result FROM answer_records')
    records = cursor.fetchall()
    conn.close()
    return render_template('student_records.html', records=records)

#学生管理
@app.route('/teacher/student_manage')
def teacher_student_manage():
    teacher = session.get('username')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 获取该教师开设的所有课程
    cursor.execute("SELECT id, name FROM courses WHERE teacher_username = ?", (teacher,))
    courses = cursor.fetchall()
    # 为每门课程查找已选学生
    course_students = []
    for course in courses:
        course_id, course_name = course
        cursor.execute('''
            SELECT users.username 
            FROM student_courses 
            JOIN users ON student_courses.student_username = users.username
            WHERE student_courses.course_id = ?
        ''', (course_id,))
        students = [row[0] for row in cursor.fetchall()]
        course_students.append((course_id, course_name, students))
    conn.close()
    return render_template('teacher_student_manage.html', course_students=course_students)


@app.route('/teacher/add_student_to_course', methods=['POST'])
def add_student_to_course():
    student_username = request.form['student_username']
    course_id = request.form['course_id']
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 验证学生是否存在
    cursor.execute("SELECT * FROM users WHERE username = ? AND role = 'student'", (student_username,))
    student = cursor.fetchone()
    if not student:
        flash("该学生不存在或身份错误！")
    else:
        # 插入选课记录（防止重复）
        try:
            cursor.execute("INSERT INTO student_courses (student_username, course_id) VALUES (?, ?)", (student_username, course_id))
            cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (session['username'], 'teacher', '添加学生', f'{student_username} -> 课程ID: {course_id}'))
            conn.commit()
            conn.close()
            flash("添加成功！")
        except sqlite3.IntegrityError:
            flash("该学生已经选过这门课了！")
    conn.close()
    return redirect(url_for('teacher_student_manage'))

@app.route('/teacher/delete_student_from_course/<int:course_id>/<student_username>')
def delete_student_from_course(course_id, student_username):
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("DELETE FROM student_courses WHERE student_username = ? AND course_id = ?", (student_username, course_id))
    cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (session['username'], 'teacher', '删除学生', f'{student_username} -> 课程ID: {course_id}'))
    conn.commit()
    conn.close()
    flash("已删除学生")
    return redirect(url_for('teacher_student_manage'))

#查看课程列表
@app.route('/teacher/courses')
def teacher_courses():
    if 'username' not in session or session.get('role') != 'teacher':
        return redirect('/login')
    teacher = session['username']
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("SELECT id, name FROM courses WHERE teacher_username = ?", (teacher,))
    courses = cursor.fetchall()
    conn.close()
    return render_template('teacher_courses.html', courses=courses)

#添加课程
@app.route('/teacher/add_course', methods=['POST'])
def add_course():
    if session.get('role') != 'teacher':
        return '无权限访问'
    teacher = session.get('username')
    course_name = request.form.get('course_name').strip()
    if not course_name:
        flash("课程名不能为空")
        return redirect('/teacher/courses')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    #判断是否已存在
    cursor.execute('SELECT * FROM courses WHERE name = ? AND teacher_username = ?', (course_name, teacher))
    existing = cursor.fetchone()
    if existing:
        flash("你已创建过该课程")
        conn.close()
        return redirect('/teacher/courses')
    # 插入课程
    cursor.execute('INSERT INTO courses (name, teacher_username) VALUES (?, ?)', (course_name, teacher))
    conn.commit()

    #自动恢复该教师该课程下曾经上传的知识库
    cursor.execute('''
        UPDATE knowledge
        SET is_active = 1
        WHERE course_name = ? AND teacher_username = ?
    ''', (course_name, teacher))
    cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (teacher, 'teacher', '添加课程', course_name))
    conn.commit()
    conn.close()
    flash("课程添加成功，相关知识库也已自动恢复")
    return redirect('/teacher/courses')

@app.route('/teacher/delete_course/<int:course_id>', methods=['POST'])
def delete_course(course_id):
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 查询课程名称（为了软删知识库）
    cursor.execute("SELECT name FROM courses WHERE id = ?", (course_id,))
    course = cursor.fetchone()
    if not course:
        conn.close()
        flash("未找到课程")
        return redirect('/teacher/courses')
    course_name = course[0]
    # 删除课程
    cursor.execute("DELETE FROM courses WHERE id = ?", (course_id,))
    # 软删知识库（隐藏）
    cursor.execute("UPDATE knowledge SET is_active = 0 WHERE course_name = ?", (course_name,))
    cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (session['username'], 'teacher', '删除课程', course_name))
    conn.commit()
    conn.close()
    flash("课程及其相关知识库已删除")
    return redirect('/teacher/courses')

#查看做题记录
@app.route('/teacher/view_records', methods=['GET'])
def teacher_view_records():
    if session.get('role') != 'teacher':
        return '无权限访问'
    teacher = session.get('username')
    course = request.args.get('course')
    student_search = request.args.get('student_search', '').strip()
    selected_type = request.args.get('type', '').strip()
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("SELECT name FROM courses WHERE teacher_username = ?", (teacher,))
    courses = [row[0] for row in cursor.fetchall()]
    students = {}
    types = ['choice', 'blank', 'short', 'code']
    if course:
        query = '''
            SELECT id, question, student_answer, correct_answer, result, student, type, feedback
            FROM answer_records
            WHERE course_name = ? AND teacher_username = ?
        '''
        params = [course, teacher]
        if student_search:
            query += " AND student LIKE ?"
            params.append(f"%{student_search}%")
        if selected_type:
            query += " AND type = ?"
            params.append(selected_type)
        cursor.execute(query, tuple(params))
        records = cursor.fetchall()
        for rec in records:
            r = {
                'id': rec[0], 'question': rec[1], 'student_answer': rec[2],
                'correct_answer': rec[3], 'result': rec[4], 'type': rec[6],
                'feedback': rec[7]
            }
            student = rec[5]
            if student not in students:
                students[student] = []
            students[student].append(r)
    all_students = list(students.keys()) if course else []
    conn.commit()
    conn.close()
    return render_template('teacher_view_records.html', courses=courses, selected_course=course,
                            student_search=student_search, types=types, selected_type=selected_type,
                            students=students,all_students=all_students)

#删除记录
@app.route('/teacher/delete_record/<int:record_id>', methods=['POST'])
def delete_answer_record(record_id):
    if session.get('role') != 'teacher':
        return '无权限访问'
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("DELETE FROM answer_records WHERE id = ?", (record_id,))
    conn.commit()
    conn.close()
    return redirect(request.referrer or url_for('teacher_view_records'))


@app.route('/teacher/generate_feedback/<int:record_id>', methods=['POST'])
def generate_feedback(record_id):
    if session.get('role') != 'teacher':
        return '无权限访问'
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute("SELECT question, student_answer, correct_answer, type FROM answer_records WHERE id = ?", (record_id,))
    rec = cursor.fetchone()
    if not rec:
        conn.close()
        flash("记录不存在")
        return redirect(request.referrer or '/teacher/view_records')
    question, student_answer, correct_answer, qtype = rec

    prompt = f"""
你是一个教育系统中的AI批改助手，请分析以下内容：
【题型】：{qtype}
【题目】：{question}
【学生回答】：{student_answer}
【标准答案】：{correct_answer}
请你对学生答案中的错误进行定位与分析，指出不足，并给出修正建议，用简洁通俗语言表达。
"""
    try:
        feedback =call_llm_for_feedback(prompt)  # 修复调用
    except Exception as e:
        feedback = f"[生成失败] {e}"
    cursor.execute("UPDATE answer_records SET feedback = ? WHERE id = ?", (feedback, record_id))
    conn.commit()
    conn.close()
    flash("错误分析已生成")
    return redirect(request.referrer or '/teacher/view_records')

@app.route('/teacher/generate_analysis', methods=['POST'])
def generate_analysis():
    if session.get('role') != 'teacher':
        return '无权限访问'
    teacher = session.get('username')
    course = request.form.get('course')
    selected_students = request.form.getlist('selected_students')
    if not course or not selected_students:
        flash("请选择课程和学生")
        return redirect('/teacher/view_records?course=' + course)
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    summary_data = []
    for student in selected_students:
        cursor.execute('''
            SELECT type, result, question, student_answer, correct_answer
            FROM answer_records
            WHERE teacher_username = ? AND course_name = ? AND student = ?
        ''', (teacher, course, student))
        records = cursor.fetchall()
        if not records:
            continue
        total = len(records)
        correct = sum(1 for r in records if r[1] >= 90)
        student_summary = {
            'student': student,
            'total': total,
            'correct': correct,
            'accuracy': round(correct / total * 100, 1),
            'type_stats': {},
            'questions': []
        }
        for r in records:
            typ = r[0]
            student_summary['type_stats'].setdefault(typ, {'total': 0, 'correct': 0})
            student_summary['type_stats'][typ]['total'] += 1
            if r[1] >= 90:
                student_summary['type_stats'][typ]['correct'] += 1
            if r[1] < 90:
                student_summary['questions'].append({
                    'question': r[2],
                    'student_answer': r[3],
                    'correct_answer': r[4]
                })
        summary_data.append(student_summary)
    conn.close()
    # 拼接 prompt
    prompt = "你是一个教学评估专家，请根据以下学生数据，逐个生成分析，最后总结整体教学建议：\n"
    for s in summary_data:
        prompt += f"\n--- 学生：{s['student']} ---\n"
        prompt += f"总题数：{s['total']}，正确数：{s['correct']}，准确率：{s['accuracy']}%\n"
        for t, d in s['type_stats'].items():
            rate = round(d['correct'] / d['total'] * 100, 1)
            prompt += f"题型 {t}：{d['correct']}/{d['total']}，正确率 {rate}%\n"
        for q in s['questions'][:2]:
            prompt += f"- 错题：{q['question']}\n  答案：{q['student_answer']}\n  正确：{q['correct_answer']}\n"

    # 分析模型接口（独立于出题）
    analysis_text = call_llm_for_analysis(prompt)
    # 保存进 session，跳转展示页
    session['analysis_result'] = analysis_text
    return redirect('/teacher/analysis_result')

@app.route('/teacher/analysis_result')
def teacher_analysis_result():
    if session.get('role') != 'teacher':
        return '无权限访问'
    result = session.get('analysis_result', '暂无分析结果')
    return render_template('teacher_analysis_report.html', analysis=result)

@app.route('/student', methods=['GET'])
def student_dashboard():
    if session.get('role') != 'student':
        return '无权限访问'
    return render_template('student.html',username=session.get('username'))

@app.route('/student/courses', methods=['GET', 'POST'])
def student_courses():
    student = session.get('username')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 获取学生已选课程
    cursor.execute('''
        SELECT c.id, c.name, c.teacher_username 
        FROM courses c 
        JOIN student_courses sc ON c.id = sc.course_id 
        WHERE sc.student_username = ?
    ''', (student,))
    selected_courses = cursor.fetchall()
    # 获取所有课程（供选择）
    cursor.execute('SELECT id, name, teacher_username FROM courses')
    all_courses = cursor.fetchall()
    conn.close()
    return render_template('student_courses.html',
                           selected_courses=selected_courses,
                           all_courses=all_courses)

@app.route('/student/choose_course', methods=['POST'])
def choose_course():
    student = session.get('username')
    course_id = request.form.get('course_id')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 检查是否已选
    cursor.execute('SELECT * FROM student_courses WHERE student_username = ? AND course_id = ?', (student, course_id))
    if cursor.fetchone():
        conn.close()
        return redirect('/student/courses')  # 已选不重复添加
    cursor.execute('INSERT INTO student_courses (student_username, course_id) VALUES (?, ?)', (student, course_id))
    cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (session['username'], 'student', '添加课程', f'课程ID: {course_id}'))#日志
    conn.commit()
    conn.close()
    return redirect('/student/courses')

@app.route('/student/delete_course/<int:course_id>')
def drop_course(course_id):
    student = session.get('username')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('DELETE FROM student_courses WHERE student_username = ? AND course_id = ?', (student, course_id))
    cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (session['username'], 'student', '退选课程', f'课程ID: {course_id}'))#日志
    conn.commit()
    conn.close()
    return redirect('/student/courses')

#选择练习页面（课程 + 题型）
@app.route('/student/practice_select')
def practice_select():
    student = session.get('username')
    if not student:
        return redirect('/')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('''
        SELECT sc.course_id, c.name, c.teacher_username
        FROM student_courses sc
        JOIN courses c ON sc.course_id = c.id
        WHERE sc.student_username = ?
    ''', (student,))
    courses = cursor.fetchall()
    conn.close()
    return render_template('practice_select.html', courses=courses)

# 路由: 开始练习（初始化练习缓存列表）
@app.route('/student/start_practice', methods=['POST'])
def start_practice():
    student = session.get('username')
    course_id = request.form.get('course_id')
    q_type = request.form.get('q_type')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('SELECT name, teacher_username FROM courses WHERE id = ?', (course_id,))
    course_info = cursor.fetchone()
    if not course_info:
        conn.close()
        flash("课程信息无效！")
        return redirect('/student/practice_select')
    cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (session['username'], 'student', '开始练习', f'课程ID: {course_id}, 题型: {q_type}')) #日志
    course_name, teacher_username = course_info
    cursor.execute('''
        SELECT * FROM question_bank
        WHERE course_name = ? AND teacher_username = ? AND type = ?
        AND id NOT IN (
            SELECT question_id FROM answer_records WHERE student = ?
        )
    ''', (course_name, teacher_username, q_type, student))
    questions = cursor.fetchall()
    conn.commit()
    conn.close()
    if not questions:
        flash("暂无该课程该类型题目或您已全部练习完成！")
        return redirect('/student/practice_select')
    session['practice_questions'] = [q[0] for q in questions]  # 存id
    session['q_type'] = q_type
    session['course_name'] = course_name
    return redirect(url_for('do_next_question'))

#学生做题
@app.route('/student/practice/next')
def do_next_question():
    q_ids = session.get('practice_questions', [])
    if not q_ids:
        flash(f"恭喜你，{session.get('course_name')}课程 {session.get('q_type')} 已完成！")
        return redirect('/student/practice_select')
    q_id = q_ids.pop(0)
    session['practice_questions'] = q_ids  # pop 一道
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM question_bank WHERE id = ?', (q_id,))
    question = cursor.fetchone()
    conn.close()
    return render_template('do_question.html', question=question, q_type=session['q_type'])

# 路由: 提交答案
@app.route('/submit_answer/<int:question_id>', methods=['POST'])
def submit_answer(question_id):
    from llm_api import evaluate_student_answer
    student = session.get('username')
    student_answer = request.form['answer']
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('SELECT question, answer, explanation, type, teacher_username, course_name FROM question_bank WHERE id = ?', (question_id,))
    question = cursor.fetchone()
    if not question:
        flash("题目不存在")
        return redirect('/student/practice_select')
    question_text, correct_answer, explanation, q_type, teacher_username, course_name = question
    if q_type in ['choice', 'blank']:
        result = 100 if student_answer.strip() == correct_answer.strip() else 0
    else:
        result = evaluate_student_answer(student_answer, correct_answer, question_text)
    cursor.execute('''
        INSERT INTO answer_records (question_id, question, student_answer, correct_answer, result, student, type, teacher_username, course_name,created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, datetime('now'))
    ''', (question_id, question_text, student_answer, correct_answer, result, student, q_type, teacher_username, course_name))
    conn.commit()
    conn.close()
    return redirect(url_for('do_next_question'))

@app.route('/student/ask', methods=['GET', 'POST'])
def student_ask():
    if session.get('role') != 'student':
        return '无权限访问'
    student = session.get('username')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 获取学生选课课程 (含课程id, 名称, 授课教师)
    cursor.execute('''
        SELECT c.id, c.name, c.teacher_username
        FROM student_courses sc
        JOIN courses c ON sc.course_id = c.id
        WHERE sc.student_username = ?
    ''', (student,))
    courses = cursor.fetchall()
    selected_course_id = request.args.get('course_id')
    selected_knowledge_list = []
    if request.method == 'GET':
        if selected_course_id:
            cursor.execute("SELECT name, teacher_username FROM courses WHERE id = ?", (selected_course_id,))
            course = cursor.fetchone()
            if course:
                course_name, teacher_username = course
                cursor.execute('''
                    SELECT id, filename FROM knowledge
                    WHERE course_name = ? AND teacher_username = ? AND is_active = 1
                ''', (course_name, teacher_username))
                selected_knowledge_list = cursor.fetchall()
        conn.commit()
        conn.close()
        return render_template(
            'student_ask.html',
            courses=courses,
            selected_course_id=selected_course_id,
            knowledge_list=selected_knowledge_list
        )
    elif request.method == 'POST':
        course_id = request.form.get('course_id')
        knowledge_id = request.form.get('knowledge_id')
        question = request.form.get('question')
        cursor.execute("SELECT name, teacher_username FROM courses WHERE id = ?", (course_id,))
        course = cursor.fetchone()
        if not course:
            return '非法课程ID'
        course_name, teacher_username = course
        if knowledge_id == 'all':
            cursor.execute('''
                SELECT content FROM knowledge
                WHERE course_name = ? AND teacher_username = ? AND is_active = 1
            ''', (course_name, teacher_username))
        else:
            cursor.execute('''
                SELECT content FROM knowledge
                WHERE id = ? AND is_active = 1
            ''', (knowledge_id,))
        knowledge_entries = cursor.fetchall()
        cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (session['username'], 'student', '学习助手问答', f'课程ID: {course_id}, 知识库ID: {knowledge_id}'))#日志
        conn.commit()
        conn.close()
        full_knowledge = "\n".join([k[0] for k in knowledge_entries])
        prompt = f"""你是一名教育AI助手，请结合以下教学知识，回答学生提出的问题。\n\n【教学知识】：\n{full_knowledge}\n\n【学生问题】：{question}\n\n请给予专业、准确、简明的解答："""
        answer = call_llm_for_ask(prompt)
        return render_template('student_ask_result.html', question=question, answer=answer)

@app.route('/student/view_records', methods=['GET'])
def view_student_records():
    student = session.get('username')
    selected_course_id = request.args.get("course_id")
    search = request.args.get("search", "")
    qtype = request.args.get("type", "")
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 获取学生选课
    cursor.execute('''
        SELECT c.id, c.name || '（' || c.teacher_username || '）'
        FROM student_courses sc
        JOIN courses c ON sc.course_id = c.id
        WHERE sc.student_username = ?
    ''', (student,))
    courses = cursor.fetchall()
    # 日志记录逻辑
    if selected_course_id and selected_course_id != "" and not session.get("view_log_done"):
        cursor.execute('''
            INSERT INTO system_logs (username, role, action, detail)
            VALUES (?, ?, ?, ?)
        ''', (student, 'student', '查看记录', f'课程ID: {selected_course_id}'))
        conn.commit()
        session["view_log_done"] = True  # 设置防重复标记
    # 正常查询答题记录逻辑
    sql = '''
        SELECT id, question, student_answer, correct_answer, result, type, teacher_username, course_name
        FROM answer_records
        WHERE student = ?
    '''
    params = [student]
    if selected_course_id:
        sql += " AND course_name = (SELECT name FROM courses WHERE id = ?)"
        params.append(selected_course_id)
    if search:
        sql += " AND question LIKE ?"
        params.append(f"%{search}%")
    if qtype:
        sql += " AND type = ?"
        params.append(qtype)
    cursor.execute(sql, tuple(params))
    records = cursor.fetchall()
    # 获取当前页所有题目 ID（为了实时练习）
    current_page_ids = [str(r[0]) for r in records]
    conn.close()
    return render_template("student_view_records.html",
                           courses=courses,
                           selected_course_id=selected_course_id,
                           search=search,
                           qtype=qtype,
                           records=records,
                           current_page_ids=current_page_ids)

def parse_generated_questions(llm_output):
    # 假设每道题用【题目】【答案】【解析】分隔
    questions = []
    parts = llm_output.strip().split("【题目】")
    for part in parts[1:]:  # 第一个是空的
        try:
            q_split = part.split("【答案】")
            question_text = q_split[0].strip()
            a_split = q_split[1].split("【解析】")
            answer = a_split[0].strip()
            explanation = a_split[1].strip()
            questions.append({
                "question": question_text,
                "answer": answer,
                "explanation": explanation
            })
        except:
            continue  # 跳过格式异常的
    return questions

@app.route('/student/realtime_practice', methods=['POST'])
def realtime_practice():
    if session.get('role') != 'student':
        return "无权限访问"
    student = session.get('username')
    course_id = request.form.get('course_id')
    filter_type = request.form.get('filter_type')
    keyword = request.form.get('keyword')
    practice_goal = request.form.get('practice_goal')
    # 接收当前页面题目的 ID 列表（前端隐藏 input 或 JS 提交）
    selected_ids = request.form.getlist("question_ids")
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    # 从数据库中查出这些题目的完整信息
    placeholder = ','.join(['?'] * len(selected_ids))
    cursor.execute(f"SELECT question, correct_answer, type FROM answer_records WHERE id IN ({placeholder})", selected_ids)
    question_info_list = cursor.fetchall()
    cursor.execute('''
    INSERT INTO system_logs (username, role, action, detail)
    VALUES (?, ?, ?, ?)
''', (session['username'], 'student', '开始实时练习', f'课程ID: {course_id}, 要求: {practice_goal}'))#日志
    conn.commit()
    conn.close()
    # 构造 prompt，调用你封装好的 call_llm_for_realtime()
    content_for_model = ""
    for q in question_info_list:
        content_for_model += f"题目：{q[0]}\n答案：{q[1]}\n题型：{q[2]}\n\n"
    prompt = f"""
注意：你是一个教育系统中的AI出题助手，请综合以下学生历史练习记录以及学生的练习要求生成类似风格的随练题目，尽可能捕捉学生曾经出错较多的题型与知识点，所生成的题目并不一定要局限于学生历史练习记录，此条注意必须严格遵守，否则我会考虑不使用deepseek！！！！！

学生历史练习记录：
{content_for_model}

学生的练习要求：
{practice_goal}

请用如下格式返回每道题：

【题目】
...

【答案】
...

【解析】
...
"""
    result = call_llm_for_realtime(prompt)
    return render_template('realtime_exercise.html', questions=parse_generated_questions(result))

@app.route('/student/evaluate_realtime_answer', methods=['POST'])
def evaluate_realtime_answer():
    if session.get('role') != 'student':
        return '无权限访问'
    student_answers = request.form.getlist("student_answers")
    correct_answers = request.form.getlist("correct_answers")
    question_texts = request.form.getlist("question_texts")
    results = []
    for sa, ca, qt in zip(student_answers, correct_answers, question_texts):
        score = evaluate_student_answer(sa, ca, qt)
        prompt = f"""请你对以下学生作答进行纠错建议：
题目：{qt}
正确答案：{ca}
学生答案：{sa}
请说明学生答案中存在的问题，并提出改进建议，不需要打分。
"""
        feedback = call_llm_for_feedback(prompt)
        results.append({
            "question": qt,
            "your_answer": sa,
            "correct_answer": ca,
            "score": score,
            "feedback": feedback
        })
    return render_template("realtime_result.html", results=results)

@app.route('/student/statistics')
def student_statistics():
    if session.get('role') != 'student':
        return '无权限访问'
    student_name = session.get('username')
    conn = sqlite3.connect('database.db')
    cursor = conn.cursor()
    cursor.execute('SELECT COUNT(*) FROM answer_records WHERE student = ?', (student_name,))
    total = cursor.fetchone()[0]
    cursor.execute('SELECT COUNT(*) FROM answer_records WHERE student = ? AND result = "正确"', (student_name,))
    correct = cursor.fetchone()[0]
    conn.close()
    accuracy = f"{(correct/total*100):.2f}%" if total else "暂无数据"
    return render_template('student_statistics.html', total=total, correct=correct, accuracy=accuracy)

if __name__ == '__main__':
    app.run(port=5000,debug=True)
